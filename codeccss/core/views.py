from django.shortcuts import render,redirect
from django.contrib.auth import login, authenticate, logout as django_logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import JsonResponse, HttpResponse
from .models import *
from core.forms import *
import json,datetime



# Create your views here.

def inicio(request):
    contexto={
        'titulo': "Inicio",
    }
    return render(request,"inicio.html",contexto)

def login_alumno(request):
    if request.POST and request.is_ajax:
        print("Matricula: "+str(request.POST.get('matricula')))
        if Expediente.objects.filter(matricula=request.POST.get('matricula')).exists():
            return HttpResponse('redirect')
        else:
            return JsonResponse({'type':'info','message':'La matricula ingresada no tiene un expediente activo'})
    contexto={
        'titulo': "Login"
    }
    return render(request,"alumno/login_alumno.html",contexto)

def login_docente(request):
    logout(request)
    if request.POST:
        user = request.POST['username']
        password = request.POST['password']

        user = authenticate(username=user, password=password)

        if user is not None:   
            login(request, user)
            if Perfil.objects.filter(user=request.user.id).exists():
                data=Perfil.objects.get(user=request.user.id)
                if data.estado:
                    return redirect("inicio_docente")
                else: 
                    return redirect("dashbord")
            else:
                usuario=User.objects.get(id=request.user.id)
                model=Perfil(user=usuario, nombre="",apellido_paterno="", apellido_materno="",email_institucional="")
                model.save() 
                return redirect("dashbord")
        else:
            messages.info(request,"Verifica el usuario y contraseña")
    else:
        return render(request,'docente/login_docente.html')

    contexto={
        'title':'Inicio de sesión',
    } 
    return render(request, 'docente/login_docente.html', contexto)


def logout(request):
    django_logout(request)
    return redirect("/")

@login_required(login_url='login_docente')
def dashbord_docente(request):
    data=Perfil.objects.get(user=request.user.id)
    form=AltaPerfilForm(instance=data)
    if request.POST:
        form=AltaPerfilForm(request.POST,request.FILES,instance=data)
        print(request.POST)
        if form.is_valid():
            form.save()
            data.estado=True
            data.save()
            print("Registro guardado correctamente")
            messages.success(request,'Perfil creado correctamente',"alert-success")
            return redirect('inicio_docente')
        else:
            print(form.errors)
            messages.warning(request,'Error durante el guardado de datos',"alert-danger")
            return redirect('dashbord')
    contexto={
        'titulo': "Dashboard",
        'usuario': Perfil.objects.get(user=request.user.id),
        'form':form
    }
    return render(request,"docente/dashbord.html",contexto)

@login_required(login_url='login_docente')
def inicio_docente(request):
    contexto={
        'titulo': "Inicio",
        'usuario': Perfil.objects.get(user=request.user.id)
    }
    return render(request,"docente/inicio.html",contexto)

@login_required(login_url='login_docente')
def editar_perfil(request):
    data=Perfil.objects.get(user=request.user.id)
    form=EditarPerfilForm(instance=data)
    url_image=""
    url_name=""
    if data.foto:
        url_image=data.foto.url
        url_name=data.foto
    if request.POST:
        form=EditarPerfilForm(request.POST,request.FILES,instance=data)
        if form.is_valid():
            form.save()
            print("Registro guardado correctamente")
            messages.success(request,'Perfil actualizado correctamente',"alert-success")
            return redirect('editar_perfil')
        else:
            messages.warning(request,'Error durante el guardado de datos',"alert-danger")
    contexto={
        'form':form,
        'titulo': "Editar Perfil",
        'usuario': Perfil.objects.get(user=request.user.id),
        'url_name': url_name,
        'url_image': url_image
    }
    return render(request,"docente/perfil.html",contexto)

@login_required(login_url='login_docente')
def expediente(request):  
    expedientes=Expediente.objects.filter(perfil__user=request.user.id,estado=True).order_by("matricula")
    print(expedientes)
    contexto={
        'titulo': "Expedientes",
        'usuario': Perfil.objects.get(user=request.user.id),
        'expedientes':expedientes,
        'total_expedientes': expedientes.count()
    }
    return render(request,"docente/expediente.html",contexto)

@login_required(login_url='login_docente')
def expedienteServicio(request):  
    expedientes=Expediente_Servicio.objects.filter(perfil__user=request.user.id,estado=True).order_by("matricula")
    print(expedientes)
    contexto={
        'titulo': "Expedientes Servicio Social",
        'usuario': Perfil.objects.get(user=request.user.id),
        'expedientes':expedientes,
        'total_expedientes': expedientes.count()
    }
    return render(request,"servicio/expediente.html",contexto)

@login_required(login_url='login_docente')
def detalle_expediente(request,id=None):
    creditos=Credito.objects.filter(user=request.user.id, matricula=id)
    modulos=Perfil.objects.get(user=request.user.id)
    contexto={
        'titulo': "Detalle Expediente",
        'modulos':modulos.modulo,
        'usuario': Perfil.objects.get(user=request.user.id),
        'creditos': creditos,
    }
    return render(request,"docente/detalle_expediente.html",contexto)

@login_required(login_url='login_docente')
def detalle_expediente_servicio(request,id=None):
    documentos=Documento_Servicio.objects.filter(user=request.user.id, matricula=id)
    modulos=Perfil.objects.get(user=request.user.id)
    contexto={
        'titulo': "Detalle Expediente Servicio Social",
        'modulos':modulos.modulo,
        'usuario': Perfil.objects.get(user=request.user.id),
        'documentos': documentos,
    }
    return render(request,"servicio/detalle_expediente.html",contexto)

@login_required(login_url='login_docente')
def Ajax_AltaExpedienteServicio(request):
    user=User.objects.get(id=request.user.id)
    perfil=Perfil.objects.get(user=user)
    form=Expediente_Servicio()
    mensaje=""
    tipo=""
    identify=0
    conteo=0
    if request.is_ajax and request.POST:
        if Expediente_Servicio.objects.filter(matricula=request.POST['matricula'], perfil=perfil).exists():
            print("Matricula ya existe no se puede agregar")
            mensaje="Expediente existente no se puede agregar"
            tipo="info"
            identify=0
        else:
            if request.POST['matricula'] != "" and request.POST['nombre'] != "" and request.POST['apellidoP'] != "" and request.POST['apellidoM'] != "":
                print("Agregando matricula")
                form.matricula=request.POST['matricula']
                form.nombre=request.POST['nombre']
                form.apellido_paterno=request.POST['apellidoP']
                form.apellido_materno=request.POST['apellidoM']
                form.perfil=perfil
                form.save()
                print("Agregado correctamente")
                mensaje="Expediente registrado correctamente"   
                tipo="success"
                conteo=Expediente_Servicio.objects.filter(perfil__user=request.user.id).count()
                print("Este es el conteo: "+str(conteo))
                identify=Expediente_Servicio.objects.last()
                identify=identify.id
            else:
                mensaje="Ha ocurrido un error verifique la informacion ingresada"
                tipo="error"
    return JsonResponse({'mensaje':mensaje,'tipo':tipo,'conteo':conteo,'id':identify})


@login_required(login_url='login_docente')
def Ajax_AltaExpediente(request):
    user=User.objects.get(id=request.user.id)
    perfil=Perfil.objects.get(user=user)
    form=Expediente()
    mensaje=""
    tipo=""
    identify=0
    conteo=0
    if request.is_ajax and request.POST:
        if Expediente.objects.filter(matricula=request.POST['matricula'], perfil=perfil).exists():
            print("Matricula ya existe no se puede agregar")
            mensaje="Expediente existente no se puede agregar"
            tipo="info"
            identify=0
        else:
            if request.POST['matricula'] != "" and request.POST['nombre'] != "" and request.POST['apellidoP'] != "" and request.POST['apellidoM'] != "":
                print("Agregando matricula")
                form.matricula=request.POST['matricula']
                form.nombre=request.POST['nombre']
                form.apellido_paterno=request.POST['apellidoP']
                form.apellido_materno=request.POST['apellidoM']
                form.perfil=perfil
                form.save()
                print("Agregado correctamente")
                mensaje="Expediente registrado correctamente"   
                tipo="success"
                conteo=Expediente.objects.filter(perfil__user=request.user.id).count()
                print("Este es el conteo: "+str(conteo))
                identify=Expediente.objects.last()
                identify=identify.id
            else:
                mensaje="Ha ocurrido un error verifique la informacion ingresada"
                tipo="error"
    return JsonResponse({'mensaje':mensaje,'tipo':tipo,'conteo':conteo,'id':identify})

@login_required(login_url='login_docente')
def Ajax_AltaDocumentoExpediente(request):
    #matricula=Expediente.objects.get(matricula=request.POST.get("matricula"))
    #model=Credito()
    user=User.objects.get(id=request.user.id)
    if request.POST and request.is_ajax:
        form=AltaCredito(request.POST, request.FILES or None)
        print(request.POST)
        if form.is_valid():
            print("Este es el documento"+str(request.FILES.get('directorio')))
            objeto=form.save(commit=False)
            objeto.nombre=request.FILES.get('directorio')
            objeto.rubro=request.POST.get('rubro')
            objeto.user=user
            #model.matricula=matricula
            #model.documento=request.FILES.get("file")
            #model.save()
            form.save()
            modelo=Credito.objects.filter(matricula=request.POST.get('matricula')).last()
            print("Registro guardado correctamente")
            return JsonResponse(data={'type':'success','message':'Documento agregado correctamente','url':str(modelo.directorio.url),'nombre':str(modelo.nombre),'rubro':str(modelo.rubro)})
        else:
            print(form.errors)
    return HttpResponse("OK")

@login_required(login_url='login_docente')
def Ajax_AltaDocumentoExpedienteServicio(request):
    user=User.objects.get(id=request.user.id)
    if request.POST and request.is_ajax:
        form=AltaDocumento(request.POST, request.FILES or None)
        print(request.POST)
        if form.is_valid():
            print("Este es el documento"+str(request.FILES.get('directorio')))
            objeto=form.save(commit=False)
            objeto.nombre=request.FILES.get('directorio')
            objeto.user=user
            form.save()
            modelo=Documento_Servicio.objects.filter(matricula=request.POST.get('matricula')).last()
            print("Registro guardado correctamente")
            return JsonResponse(data={'type':'success','message':'Documento agregado correctamente','url':str(modelo.directorio.url),'nombre':str(modelo.nombre)})
        else:
            print(form.errors)
    return HttpResponse("OK")

def expedienteAlumno(request,matricula=None):    
    print(str(matricula))
    data=Credito.objects.filter(matricula__matricula=matricula)
    documentos=Documento_Servicio.objects.filter(matricula__matricula=matricula)
    print("Estos son los creditos: "+str(data))
    print("Estos son los documentos: "+str(documentos))

    contexto={
        'title':'Busqueda Expediente',
        'exp':data,
        'documentos':documentos,
    }
    return render(request,"alumno/detalle_expediente.html",contexto)

@login_required(login_url='login_docente')
def historial(request):
    option=Modulo.objects.all()
    contexto={
        'option':option,
        'usuario': Perfil.objects.get(user=request.user.id),
        'title':'Historial',
        'titulo': "Historial",
    }
    return render(request,'docente/historial.html',contexto)

@login_required(login_url='login_docente')
def modulo(request):
    if request.is_ajax and request.POST:
        if request.POST['tipo'] == "add":
            modelo=Modulo()
            modelo.tipo=request.POST['modulo']
            modelo.maximo=request.POST['valor']
            modelo.save()
            info=Modulo.objects.get(tipo=request.POST['modulo'])
            return JsonResponse({'id':info.id})
        elif request.POST['tipo'] == "update":
            Modulo.objects.filter(id=request.POST['id']).delete()
    contexto={
        'title':'Modulo',
        'titulo':'Modulo',
        'usuario': Perfil.objects.get(user=request.user.id),
        'data':Modulo.objects.all(),
    }
    return render(request,'docente/modulo.html',contexto)

@login_required(login_url='login_docente')
def alta_credito(request):
    plantillas=None
    modulos=Perfil.objects.get(user=request.user.id)
    if Plantilla.objects.filter(user=request.user.id).exists():
        plantillas=Plantilla.objects.filter(user=request.user.id)
    contexto={
        'title':'Redactar credito',
        'titulo':'Redactar credito',
        'modulos':modulos.modulo,
        'usuario': Perfil.objects.get(user=request.user.id),
        'plantillas':plantillas,
    }
    return render(request,'docente/alta_credito.html',contexto)

@login_required(login_url='login_docente')
def Ajax_ConsultaHistorial(request):
    if request.is_ajax and request.POST:
        tiempo=datetime.datetime.strptime(request.POST.get('date_end'), "%Y-%m-%d").date()
        dia=tiempo+datetime.timedelta(days=1)
        print(dia)
        array=json.loads(request.POST.get('module'))
        data=Expediente.objects.filter(perfil__modulo__tipo__in=array).filter(created_at__range=[str(request.POST.get('date_start')),str(dia)])
        
        print("Esta es la data "+str(data))
    return HttpResponse("ok")

from docx import Document
import locale
#from django.core.files import File as DjangoFile
from os import remove
import shutil
from docx2pdf import convert
import pythoncom
from django.db.models import Sum
from django.conf import settings

@login_required(login_url='login_docente')
def Ajax_RedactarCredito(request):
    if request.is_ajax and request.POST:
        pythoncom.CoInitialize()
        conteo={}
        tipo=""
        mensaje=""
        doc=""
        nombre_doc=request.POST.get("matricula")+"_"+request.POST.get("nombre")
        if Plantilla.objects.filter(id=request.POST.get("select_docplantilla")).exists():
            data=Plantilla.objects.get(id=request.POST.get("select_docplantilla"))
            doc=Document(data.archivo.path)
            print("Este es el doc "+str(doc))
            print("Esta entrando en la plantilla predeterminada")
        else:
            print("Esta entrando en la plantilla definida")
            print("Este es el doc "+str(doc))
            doc=Document(request.FILES.get("documento"))
        
        if Historial_Redaccion.objects.filter(matricula=request.POST.get("matricula"), modulo=request.POST.get("modulo")).exists():
            conteo=Historial_Redaccion.objects.filter(matricula=request.POST.get("matricula"), modulo=request.POST.get("modulo")).aggregate(Sum('valor'))
        else:
            conteo={"valor__sum":0}
        data=Modulo.objects.get(tipo=str(request.POST.get("modulo")))
        print("Este es el conteo: "+str(conteo)+" y el limite: "+str(data.maximo))
        diferencia=int(data.maximo)-int(conteo.get("valor__sum"))
        print("--------- Conteo Historial ---------")
        print("Este es el valor de los creditos asignados: "+str(conteo))
        print("------------------------------------")
        if diferencia >= int(request.POST.get("valor")):
            today = datetime.datetime.now()
            locale.setlocale(locale.LC_ALL, ("es_ES", "UTF-8"))
            data=Perfil.objects.get(user=request.user)
            inline=""
            dict={"ENCARGADO":str(data.nombre+" "+data.apellido_paterno+" "+data.apellido_materno),"ALUMNO":request.POST.get("nombre"),"MATRICULA":request.POST.get("matricula"),
            "CARRERA":request.POST.get("carrera"),"RUBRO":request.POST.get("modulo"),"PERIODO":request.POST.get("periodo"),
            "VALOR":request.POST.get("valor"),"DIA":today.strftime("%d"),"MES":today.strftime("%B"),"AÑO":today.strftime("%Y"),"JEFA":request.POST.get("jefa")}
            for values in dict.keys():
                for p in doc.paragraphs:
                    if values in p.text:
                        inline = p.runs
                    # Loop added to work with runs (strings with same style)
                    for i in range(len(inline)):
                        if values in inline[i].text:
                            text = inline[i].text.replace(values, dict[values])
                            inline[i].text = text
            doc.save(nombre_doc+".docx")
            print("Se ha guardado la plantilla")
            # Change word to pdf
            convert(nombre_doc+".docx")
            convert(nombre_doc+".docx",nombre_doc+".pdf")

            remove(nombre_doc+".docx")
            os.rename(nombre_doc+".pdf",nombre_doc+".pdf")
            desktop = os.path.join(os.path.join(os.path.expanduser('~')), 'Desktop') 
            if os.path.isfile(desktop+"/"+nombre_doc+".pdf"):
                shutil.copy(nombre_doc+".pdf",desktop)
                remove(nombre_doc+".pdf")
            else:
                shutil.move(nombre_doc+".pdf",desktop)
            
            ##Guardado de historial 
            Historial_Redaccion.objects.create(matricula=request.POST.get('matricula'),nombre=request.POST.get("nombre"),carrera=request.POST.get("carrera"),modulo=request.POST.get("modulo"),valor=request.POST.get("valor"),jefa=request.POST.get('jefa'))
            tipo="success"
            mensaje="realizado"

        #file_obj1 = DjangoFile(open("Plantilla.docx", mode='rb'), name="Plantilla.docx")
        #Plantilla.objects.create(archivo=file_obj1,user=request.user)
        else:
            print("Esta entrando en el else porque revaso los limites")
            tipo="info"
            mensaje="denegado"
    return JsonResponse({'type':tipo,'result':mensaje})

@login_required(login_url='login_docente')
def Ajax_GuardarHistorialCredito(request):
    if request.is_ajax and request.GET:
        if Historial_Redaccion.objects.filter(matricula=request.GET.get('matricula')).exists():
            data=Historial_Redaccion.objects.filter(matricula=request.GET.get('matricula')).last()
            return JsonResponse({'nombre':str(data.nombre),'carrera':str(data.carrera),'jefa':str(data.jefa)})
        else:
            return HttpResponse("info")
    return HttpResponse("ok")

@login_required(login_url='login_docente')
def Ajax_GuardarPlantilla(request):
    if request.POST and request.is_ajax:
        print("Entra al metodo para guardar la plantilla")
        usuario=User.objects.get(id=request.user.id)
        Plantilla.objects.create(archivo=request.FILES.get("plantilla"),user=usuario)
        print("Se ha guardado correctamente")
    return HttpResponse("OK")